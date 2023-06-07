# ==========================чтение документа файла, полученного из шаблона ()===========================================
# значения переменных читаются из файла и записываются в БД PostgreSQL
import os
from sys import argv, exit
import docx
import psycopg2
import openpyxl

try:
  prname, file_name,file_templ,WorkDoc,WorkBPR,WorkBP,WorkENTP = argv    # соответствующий файлу шаблон
  varvel = [] # данные переменной из БД [переменная,значение,тип ]

#==========================читаем из БД значение переменной=============================================================
  def get_var_val(s):# считать атрибуты переменной из БД s - имя переменной
     ss = f"'{s}'"
     sa = 'select docp_p,docp_v,docp_t,docp_r,docp_c from docp where docp_p='+ss
     cursor.execute(sa)
     return cursor.fetchone()
  def save_bd(s):# записать значение переменной в БД
     ss = f"'{s}'"
     sss = f"'{varvel[0]}'"
     wd = str(WorkDoc)
     wr = str(WorkBPR)
     wp = str(WorkBP)
     we = str(WorkENTP)
     if varvel[0][2] == 'E':
         wd = '0'
         wr = '0'
         wp = '0'
     elif varvel[0][2] == 'P':
         wd = '0'
         wr = '0'
     elif varvel[0][2] == 'R':
         wd = '0'
     sa = 'update docp set docp_v='+ss+' where docp_p='+sss+' and doc_id ='+wd+' and bpr_id ='+wr+ \
          ' and bp_id =' + wp +' and entp_id ='+we
     cursor.execute(sa)

#============================находим переменную ее значение и координаты в таблице======================================
  def var_find(s):
      m = 0
      n = 0
      global varvel, cell
      while m != -1:
          m = s.find('${', n)
          if m != -1:
              n = s.find('}', m + 1)
              if n == -1:
                print(' ошибка в шаблоне ${..')
                return  -1
              else:
                ss = s[m:n + 1]           # имя переменной
                varvel = get_var_val(ss)  # находим ее атрибуты
                if  (varvel[1] != None):
                  print('значение переменной '+ss+'  задано')
                  return -1
                return 0
          else:
              return -1
#============================чтение  документа  по шаблону=============================================================
  def doc_read(file_name,file_templ):
#------------------------------обработка файлов txt---------------------------------------------------------------------
      if ext[1] in ['.txt', '.html']:
        with open(file_name, 'r', encoding='UTF-8') as fi,\
             open(file_templ, 'r', encoding='UTF-8') as fw:
          lines = fi.readlines()

#-------------------------------python-docx----------------------------------------------------------------------------
      if ext[1] in ['.odt','odf','odp','.docx']:
         doc = docx.Document(file_name)   # заполненный документ
         tmp = docx.Document(file_templ)  # шаблон документа
         if len(tmp.paragraphs) > 1:
            for par in tmp.paragraphs:
                if var_find(par.text) == 0:
                   print(' найти значение переменной невозможно')
                   return -1
         n = len(tmp.tables)
         if n > 0:
            n = -1
            for tab in tmp.tables:
                n += 1     # номер таблицы
                for ro in tab.rows:
                    for cell in ro.cells:
                        if len(cell.text) > 4:
                           if var_find(cell.text) == 0: # находим переменную в шаблоне
                              table = doc.tables[n]
                              s = table.cell(varvel[3]-1,varvel[4]-1).text
                              save_bd(s)

#------------------------------------openpyxl---------------------------------------------------------------------------
      if ext[1] =='.xlsx':
         workbook = openpyxl.load_workbook(file_name,data_only=True)
         sheet = workbook.active
         for row in range(sheet.max_row):
             row += 1
             for sel in range(sheet.max_column):
                 sel += 1
                 val = sheet.cell(row, sel).value
                 if val != None:
                   if (len(val) > 4):
                      if var_find(val) == 0:
                         sheet.cell(row,sel).value = sheet.cell(row,sel).value.replace(varvel[0],varvel[1])

 #======================основная программа=============================================================================
  conn = psycopg2.connect(host='localhost', database='BP', user='postgres', password='rfn15')
  # Получаем объект курсора для выполнения SQL-запросов
  cursor = conn.cursor()
  conn.autocommit = True
  ext = os.path.splitext(file_name)
  doc_read(file_name,file_templ)
  cursor.close()
  conn.close()
  exit(0)

except FileNotFoundError:
     print('file not found-' + file_name)
     exit(-1)