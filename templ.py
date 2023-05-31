# выделение из шаблона документа переменных замены, их описания, координаты в таблице -> запись в БД таблицы DOCP
#
import os
import re
from sys import argv, exit
import docx
import openpyxl
import psycopg2
try:
  prname, file_name,WorkDoc,WorkBPR,WorkBP,WorkENTP = argv
  # file_name,  файл шаблона
  # WorkDoc,    документ            ${C....} переменная конкретного документа, если WorkDoc != 0
  # WorkBPR,    экземпляр процесса  $(R....} переменная всех документов конкретного процесса, если WorkDoc =0 & WorkBPR != 0
  # WorkBP,     процесс             ${P....} переменная всех экземляров процесса. ----------\\----------------------------
  # WorkENTP    предприятие         ${E....} переменная всех процессов предприятия ----------\\--------------------------

  ss = ''
  sss = ''
  pr_cell = ''
  bef_cell = '' # текст перед переменной
  aft_cell = '' # текс после переменной
  row = 0  # номер строки переменной замены
  sel = 0  # номер ячейки в строке

#=================поиск переменных в строке и запись их в БД============================================
  def ParAdd(s,doc=False):
     global WorkDoc, row, sel, bef_cell, aft_cell
     n = 0
     j = 0
     m = 0
     ret = -1
     while m != -1:
        m = s.find('${', n)
        if m != -1:
            n = s.find('}', m + 1)
            if n == -1:
                print(' ошибка в шаблоне ${..')
                exit
            ss = "'"+s[m:n + 1]+"'"  # ключ
            wd = str(WorkDoc)
            wr = str(WorkBPR)
            wp = str(WorkBP)
            we = str(WorkENTP)
            typ = 'string'
            if doc:
               sss = ' '
               if aft_cell != '':
                  sss = aft_cell
               else:
                   if bef_cell != '':
                      sss =bef_cell
            else:
                sss = s[j: m-1]
            if (re.search(r"[а-яА-Я]",sss) == None) or (sss.find('${') != -1) :
               sss = ' '
            if s[2] == 'd':
               sss = "текущий день"
            elif s[2] =='m':
                sss = "текущий месяц"
            elif s[2] =='y':
                sss = "текущий год"
            elif s[2] =='s':
                sss = "serial"
                typ = 'serial'
            elif s[2] == 'D':
               sss = "день"
            elif s[2] =='M':
                sss = "месяц"
            elif s[2] =='Y':
                sss = "год"
            elif s[2] == 'E':
                wd =f"'0'"
                wr = f"'0'"
                wp = f"'0'"
            elif s[2] == 'P':
                wd =f"'0'"
                wr = f"'0'"
            elif s[2] == 'R':
                wd =f"'0'"

            s = f"'{sss}'"
            sa = 'insert into docp (docp_p,docp_d,docp_t,docp_r,docp_c,doc_id,bpr_id,bp_id,entp_id) values ('\
                 + ss + ',' + s + ',' + typ+','+str(row)+','+str(sel)+','+ wd+','+wr+','+wp+','+we+') on conflict do nothing'
            cursor.execute(sa)
            ret = 0
        j = n + 1
     return ret
  def GetPL(file_name):
      global WorkDoc, row, sel, bef_cell, aft_cell

      ext = os.path.splitext(file_name)
#=================обработка файлов txt=======================================
      if ext[1] in ['.txt', '.html']:
        with open(file_name, 'r', encoding='UTF-8') as fi:
          lines = fi.readlines()
          row = 0
          sel = 0
          for line in lines:
              row += 1
              ParAdd(line)
#=======================python-docx=======================================================
      if ext[1] in ['.doc','.docx']:
         doc = docx.Document(file_name)
#===================абзацы===============================================================
         if len(doc.paragraphs) > 1:
            row = 0
            for par in doc.paragraphs:
                row += 1
                ParAdd(par.text)
#==================таблицы - шаблоны документов==========================================
         n = len(doc.tables)    # кол.таблиц
         if n > 0: #
           for tab in doc.tables:
                n = len(tab.rows)
                row = 0
                for ro in tab.rows:
                   row += 1
#                   n = len(ro.cells)
                   bef_cell = ''
                   aft_cell = ''
                   sel = 0
                   for cell in ro.cells:
                       sel += 1
                       s = cell.text
                       if (len(s) > 3) and (s != bef_cell):   # not in [pr_cell,'.',',']:
                          if row < n:
                            ce = tab.cell(row,sel)
                            if ce != None:
                               aft_cell = ce.text
                          ParAdd(s,True)
                          bef_cell = s

#=======================xlrd=======================================================
      if ext[1] =='.xls':
         workbook = xlrd.open_workbook(file_name)
         sheets_name = workbook.sheet_names()
         for names in sheets_name:
             worksheet = workbook.sheet_by_name(names)
             num_rows = worksheet.nrows
             num_cells = worksheet.ncols
             for row in range(num_rows):
                 pr_cell = ''
                 row += 1
                 for sel in range(num_cells):
                     val = worksheet.cell_value(row, sel)
                     sel += 1
                     if val != None:
                       if (len(val) > 2) and (val != pr_cell):
                          ParAdd(val)
                     #     print(val)
                          pr_cell = val
#=============================openpyxl==================================================
      if ext[1] =='.xlsx':
         workbook = openpyxl.load_workbook(file_name,data_only=True)
         sheet = workbook.active
         for row in range(sheet.max_row):
             pr_cell = ''
             row += 1
             for sel in range(sheet.max_column):
                 sel += 1
                 val = sheet.cell(row, sel).value
                 if val != None:
                   if (len(val) > 2) and (val != pr_cell):
                      ParAdd(val)
                    #  print(val)
                      pr_cell = val

#  cursor = DBEngine('BP')
  conn = psycopg2.connect(host='localhost', database='BP', user='postgres', password='rfn15')
# Получаем объект курсора для выполнения SQL-запросов
  cursor = conn.cursor()
  conn.autocommit = True
  GetPL(file_name)
  cursor.close()
  conn.close()
  exit(0)
except FileNotFoundError:
     print('file not found-' + file_name)
     exit(-1)